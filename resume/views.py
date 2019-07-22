from django.shortcuts import render, redirect
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.forms import User
from django.views.generic import TemplateView
from docx import Document
from docx.shared import Pt
from .forms import RegistrationForm, LoginForm, InputForm1, InputForm2, InputForm3, InputForm4
from django.contrib.auth.mixins import LoginRequiredMixin
from .models import Resumedata


def home(request):
    context = {
        'message': '英文履歴書作成サイトへようこそ！',
        'steps': ['Step1', 'Step2', 'Step3'],
        'input_types': ['個人情報', '経歴', 'あなたについて']
    }
    return render(request, 'resume/home.html', context)



class Register(TemplateView):
    template_name = 'resume/register.html'

    def get(self, request):
        form = RegistrationForm()

        context = {'Form': form, 'message': 'ユーザー登録をしてください', 'button': '登録'}
        return render(request, self.template_name, context)

    def post(self, request):
        form = RegistrationForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('step1')

        else:
            form = RegistrationForm()
#       no-valid error message to be added.
        return render(request, self.template_name, {'Form': form} )



class Login(TemplateView):
    template_name = 'resume/register.html'

    def get(self, request):
        form = LoginForm()

        context = {'Form': form, 'message': 'ログインしてください', 'button': 'ログイン'}
        return render(request, self.template_name, context)

    def post(self, request):
        form = LoginForm(request.POST)
        username = request.POST['username']
        password = request.POST['password']
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            return redirect('step1')
        else:
            form = LoginForm()
            return render(request, self.template_name, {'Form': form} )



def logout_view(request):
    logout(request)
    return redirect('home')


class InputView(LoginRequiredMixin, TemplateView):
    template_name = 'resume/step.html'

    def get(self, request):
        inputForm = InputForm1()

        context = {'inputForm': inputForm, 'message': '連絡先・アカウント情報を入力してください'}
        return render(request, self.template_name, context)

    def post(self, request):
        inputForm = InputForm1(request.POST, instance=request.user.resumedata)
        if inputForm.is_valid():
            data = inputForm.save(commit=False)
#            post.user = request.user
            data.save()
            return redirect('step2')

        return render(request, self.template_name)



class InputView2(LoginRequiredMixin, TemplateView):
    template_name = 'resume/step.html'

    def get(self, request):
        inputForm = InputForm2()

        context = {'inputForm': inputForm, 'message': '学校について教えてください'}
        return render(request, self.template_name, context)

    def post(self, request):
        inputForm = InputForm2(request.POST, instance=request.user.resumedata)
        if inputForm.is_valid():
            data = inputForm.save(commit=False)
            data.save()
#            school1 = inputForm.cleaned_data['school1']
#            major1 = inputForm.cleaned_data['major1']
#            s_start1 = inputForm.cleaned_data['s_start1']
#            s_end1 = inputForm.cleaned_data['s_end1']
#            school2 = inputForm.cleaned_data['school2']
#            major2 = inputForm.cleaned_data['major2']
#            s_start2 = inputForm.cleaned_data['s_start2']
#            s_end2 = inputForm.cleaned_data['s_end2']
#            school3 = inputForm.cleaned_data['school3']
#            major3 = inputForm.cleaned_data['major3']
#            s_start3 = inputForm.cleaned_data['s_start3']
#            s_end3 = inputForm.cleaned_data['s_end3']
            return redirect('step3')

        context = {
            'school1': school1,
            'major1': major1,
            's_start1': s_start1,
            's_end1': s_end1,
            'school2': school2,
            'major2': major2,
            's_start2': s_start2,
            's_end2': s_end2,
            'school3': school3,
            'major3': major3,
            's_start3': s_start3,
            's_end3': s_end3,
        }
        return render(request, self.template_name, context)



class InputView3(LoginRequiredMixin, TemplateView):
    template_name = 'resume/step3.html'

    def get(self, request):
        inputForm = InputForm3()

        context = {'inputForm': inputForm, 'message': '自己PRを作成しましょう',
                   'aboutme1': ' [希望職種] ', 'aboutme2': ' [付けたいスキル] '}
        return render(request, self.template_name, context)

    def post(self, request):
        inputForm = InputForm3(request.POST, instance=request.user.resumedata)
        if inputForm.is_valid():
            data = inputForm.save(commit=False)
            data.save()
            aboutme1 = inputForm.cleaned_data['aboutme1']
            aboutme2 = inputForm.cleaned_data['aboutme2']
            #return redirect('step4')

        context = {
            'inputForm': inputForm,
            'message': '自己PRを作成しましょう',
            'aboutme1': aboutme1,
            'aboutme2': aboutme2,
        }
        return render(request, self.template_name, context)



class InputView4(LoginRequiredMixin, TemplateView):
    template_name = 'resume/step.html'

    def get(self, request):
        inputForm = InputForm4()

        context = {'inputForm': inputForm, 'message': '職歴・スキルをアピールしましょう'}
        return render(request, self.template_name, context)

    def post(self, request):
        inputForm = InputForm4(request.POST, instance=request.user.resumedata)
        if inputForm.is_valid():
            data = inputForm.save(commit=False)
            data.save()
            return redirect('profile')

#            exp1 = inputForm.cleaned_data['exp1']
#            position1 = inputForm.cleaned_data['position1']
#            e_start1 = inputForm.cleaned_data['e_start1']
#            e_end1 = inputForm.cleaned_data['e_end1']
#            exp2 = inputForm.cleaned_data['exp2']
#            position2 = inputForm.cleaned_data['position2']
#            e_start2 = inputForm.cleaned_data['e_start2']
#            e_end2 = inputForm.cleaned_data['e_end2']
#            exp3 = inputForm.cleaned_data['exp3']
#            position3 = inputForm.cleaned_data['position3']
#            e_start3 = inputForm.cleaned_data['e_start3']
#            e_end3 = inputForm.cleaned_data['e_end3']
#            skill_type1 = inputForm.cleaned_data['skill_type1']
#            skill_level1 = inputForm.cleaned_data['skill_level1']
#            skill_type2 = inputForm.cleaned_data['skill_type2']
#            skill_level2 = inputForm.cleaned_data['skill_level2']
#            skill_type3 = inputForm.cleaned_data['skill_type3']
#            skill_level3 = inputForm.cleaned_data['skill_level3']



#           return CreateResume(request)


#        context = {
#            'exp1': exp1,
#            'position1': position1,
#            'e_start1': e_start1,
#            'e_end1': e_end1,
#            'exp2': exp2,
#            'position2': position2,
#            'e_start2': e_start2,
#            'e_end2': e_end2,
#            'exp3': exp3,
#            'position3': position3,
#            'e_start3': e_start3,
#            'e_end3': e_end3,
#            'skill+type1': skill_type1,
#            'skill_level1': skill_level1,
#            'skill+type2': skill_type2,
#            'skill_level2': skill_level2,
#            'skill+type3': skill_type3,
#            'skill_level3': skill_level3,

#        }
        return render(request, self.template_name, context)


def user_profile(request):
    context = {'user': request.user, 'message': '内容を確認してください'}
    return render(request, 'resume/profile.html', context)



def combine(first, second):
    return str(first) + " - " + str(second)


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


class Myheading1:
    def __init__(self, run):
        font = run.font
        font.name = "Calibri"
        font.size = Pt(26)
        font.bold = True


class Myheading2:
    def __init__(self, run):
        font = run.font
        font.name = "Calibri"
        font.size = Pt(16)


class Mybody:
    def __init__(self, run):
        font = run.font
        font.name = "Calibri"
        font.size = Pt(12)


def createresume(request):
    document = Document('template.docx')
    user = request.user
    resumeinput = request.user.resumedata
    name = str(user.first_name) + " " + str(user.last_name)
    s_period1 = combine(resumeinput.s_start1, resumeinput.s_end1)
    s_period2 = combine(resumeinput.s_start2, resumeinput.s_end2)
    s_period3 = combine(resumeinput.s_start3, resumeinput.s_end3)
    e_period1 = combine(resumeinput.e_start1, resumeinput.e_end1)
    e_period2 = combine(resumeinput.e_start2, resumeinput.e_end2)
    e_period3 = combine(resumeinput.e_start3, resumeinput.e_end3)
    skill1 = combine(resumeinput.skill_type1, resumeinput.skill_level1)
    skill2 = combine(resumeinput.skill_type2, resumeinput.skill_level2)
    skill3 = combine(resumeinput.skill_type3, resumeinput.skill_level3)

    para1 = document.paragraphs[1]
    run1 = para1.add_run(name)
    run1 = Myheading1(run1)

    para3 = document.paragraphs[3]
    run3 = para3.add_run(resumeinput.phone)
    run3 = Mybody(run3)

    para4 = document.paragraphs[4]
    run4 = para4.add_run(user.email)
    run4 = Mybody(run4)

    para5 = document.paragraphs[5]
    run5 = para5.add_run(resumeinput.twitter)
    run5 = Mybody(run5)

    para6 = document.paragraphs[6]
    run6 = para6.add_run(resumeinput.instagram)
    run6 = Mybody(run6)

    para7 = document.paragraphs[7]
    run7 = para7.add_run(resumeinput.facebook)
    run7 = Mybody(run7)

    para9 = document.paragraphs[9]
    run9 = para9.add_run(resumeinput.school1)
    run9 = Myheading2(run9)

    para10 = document.paragraphs[10]
    run10 = para10.add_run(resumeinput.major1)
    run10 = Mybody(run10)

    para11 = document.paragraphs[11]
    run11 = para11.add_run(s_period1)
    run11 = Mybody(run11)

    para12 = document.paragraphs[12]
    run12 = para12.add_run(resumeinput.school2)
    run12 = Myheading2(run12)

    para13 = document.paragraphs[13]
    run13 = para13.add_run(resumeinput.major2)
    run13 = Mybody(run13)

    para14 = document.paragraphs[14]
    run14 = para14.add_run(s_period2)
    run14 = Mybody(run14)

    para15 = document.paragraphs[15]
    run15 = para15.add_run(resumeinput.school3)
    run15 = Myheading2(run15)

    para16 = document.paragraphs[16]
    run16 = para16.add_run(resumeinput.major3)
    run16 = Mybody(run16)

    para17 = document.paragraphs[17]
    run17 = para17.add_run(s_period3)
    run17 = Mybody(run17)

    para24 = document.paragraphs[24]
    run24 = para24.add_run(resumeinput.exp1)
    run24 = Myheading2(run24)

    para25 = document.paragraphs[25]
    run25 = para25.add_run(resumeinput.position1)
    run25 = Mybody(run25)

    para26 = document.paragraphs[26]
    run26 = para26.add_run(e_period1)
    run26 = Mybody(run26)

    para27 = document.paragraphs[27]
    run27 = para27.add_run(resumeinput.exp2)
    run27 = Myheading2(run27)

    para28 = document.paragraphs[28]
    run28 = para28.add_run(resumeinput.position2)
    run28 = Mybody(run28)

    para29 = document.paragraphs[29]
    run29 = para29.add_run(e_period2)
    run29 = Mybody(run29)

    para30 = document.paragraphs[30]
    run30 = para30.add_run(resumeinput.exp3)
    run30 = Myheading2(run30)

    para31 = document.paragraphs[31]
    run31 = para31.add_run(resumeinput.position3)
    run31 = Mybody(run31)

    para32 = document.paragraphs[32]
    run32 = para32.add_run(e_period3)
    run32 = Mybody(run32)

    para34 = document.paragraphs[34]
    run34 = para34.add_run(skill1)
    run34 = Myheading2(run34)

    para35 = document.paragraphs[35]
    run35 = para35.add_run(skill2)
    run35 = Myheading2(run35)

    para36 = document.paragraphs[36]
    run36 = para36.add_run(skill3)
    run36 = Myheading2(run36)

    para37 = document.paragraphs[37]
    para37.clear()
    len(para37.text) == 0
    delete_paragraph(para37)

    para37 = document.paragraphs[37]
    para37.clear()
    len(para37.text) == 0
    delete_paragraph(para37)

    filename = (name+'.docx')
    filepath = r'/Users/taka/Documents/Projects/myapp/resume/static/resume/' + filename
    document.save(filepath)

    import cloudconvert

    api = cloudconvert.Api('D7z0qdhAi6bRFoMXj2Yncay1QIPFkAbHBNPM0JFGufHa25KVx7amOwJAiJ7qhaHD')
    process = api.convert({
        'inputformat': 'docx',
        'outputformat': 'png',
        'input': 'upload',
        'file': open(filename, 'rb')
    })
    process.wait()
    pngname = (name+'.png')
    pngpath = r'/Users/taka/Documents/Projects/myapp/resume/static/resume/' + pngname
    process.download(pngpath)

    return redirect('draft')

def draft(request):
    pngname = (request.user.first_name + request.user.last_name)
    context = {
        'message': 'これでいいですか？',
        'pngname': pngname
    }
    return render(request, 'resume/draft.html', context)